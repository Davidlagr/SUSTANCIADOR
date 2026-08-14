import streamlit as st
import docx
from docx import Document
import io
import pandas as pd
from datetime import datetime, date

st.set_page_config(page_title="Asistente de Resoluciones - Liquidador Pensional Pro", layout="wide")

st.title("⚖️ Asistente de Sustanciación de Resoluciones")
st.subheader("Módulo de Actos Administrativos - Prestaciones Económicas")

# 1. Datos del Asegurado / Solicitante
st.header("1. Datos del Asegurado / Solicitante")
col1, col2 = st.columns(2)
with col1:
    nombres = st.text_input("Nombres")
    apellidos = st.text_input("Apellidos")
with col2:
    identificacion = st.text_input("Número de Identificación")
    # Se ajusta el calendario para permitir fechas desde 1930
    nacimiento = st.date_input(
        "Fecha de Nacimiento", 
        min_value=date(1930, 1, 1), 
        max_value=date.today(),
        value=date(1960, 1, 1),
        format="DD/MM/YYYY"
    )

# 2. Antecedentes
st.header("2. Antecedentes")
antecedentes = st.text_area("Pegue aquí los antecedentes del trámite (Radicados, historias laborales, solicitudes previas):", height=150)

# 3. Banco de Motivaciones (Consideraciones Legales)
st.header("3. Banco de Motivaciones Legales")
banco_motivaciones = {
    "Sustitución Pensional - Marco General (Ley 797/2003)": "Que el artículo 47 de la citada Ley 100 de 1993, modificado por el artículo 13 de la Ley 797 de 2003 establece como beneficiarios de la pensión de sobrevivientes: a) En forma vitalicia, el cónyuge o la compañera o compañero permanente supérstite, siempre y cuando dicho beneficiario, a la fecha del fallecimiento del causante, tenga 30 o más años de edad...",
    "Sustitución Pensional - Hijo Inválido (Dependencia)": "Que frente a la dependencia económica de los hijos inválidos, el Memorando OAL-001-2022 del 13 de enero de 2022 emitido por la Oficina Asesora de Asuntos Legales, establece que para acreditar la dependencia económica del hijo inválido NO se requiere probar la carencia total y absoluta de medios económicos, existiendo subordinación cuando la persona requiera total o parcialmente de los ingresos de otra para cubrir sus necesidades básicas...",
    "Pensión de Vejez - Fórmula Decreciente": "Que para establecer el monto de la pensión, se aplicará la fórmula decreciente estipulada en el artículo 10 de la Ley 797 de 2003, que modificó el artículo 34 de la Ley 100 de 1993, determinando la tasa de reemplazo según el número de salarios mínimos del IBL y las semanas adicionales cotizadas...",
    "Aplicación de IPC (Mesada)": "Que el valor de la mesada será reajustado al momento del pago, según el Índice de Precios al Consumidor certificado por el DANE, de acuerdo con lo establecido por el artículo 14 de la Ley 100 de 1993."
}

motivaciones_seleccionadas = st.multiselect(
    "Seleccione las motivaciones a incluir en las CONSIDERACIONES:",
    options=list(banco_motivaciones.keys())
)

texto_motivaciones = "\n\n".join([banco_motivaciones[m] for m in motivaciones_seleccionadas])
if texto_motivaciones:
    st.info("Motivaciones seleccionadas preparadas para el acto administrativo.")

# 4. Fórmula Decreciente (Ley 797 de 2003)
st.header("4. Cálculo y Explicación de Fórmula Decreciente (Ley 797)")

# BASE DE DATOS SMLMV - ACTUALIZACIÓN ANUAL DESDE 1984
smlmv_historico = {
    2026: 1750905,
    2025: 1423500,
    2024: 1300000, 2023: 1160000, 2022: 1000000, 2021: 908526,
    2020: 877803, 2019: 828116, 2018: 781242, 2017: 737717,
    2016: 689455, 2015: 644350, 2014: 616000, 2013: 589500,
    2012: 566700, 2011: 535600, 2010: 515000, 2009: 496900,
    2008: 461500, 2007: 433700, 2006: 408000, 2005: 381500,
    2004: 358000, 2003: 332000, 2002: 309000, 2001: 286000,
    2000: 260100, 1999: 236460, 1998: 203826, 1997: 172005,
    1996: 142125, 1995: 118934, 1994: 98700, 1993: 81510,
    1992: 65190, 1991: 51716, 1990: 41025, 1989: 32560,
    1988: 25637, 1987: 20510, 1986: 16811, 1985: 13558,
    1984: 11298
}

current_year = datetime.now().year

col3, col4, col5 = st.columns(3)
with col3:
    semanas_totales = st.number_input("Total Semanas Cotizadas", min_value=1300.0, value=1300.0, step=1.0)
    semanas_adicionales = int(max(0, semanas_totales - 1300))
    st.caption(f"Semanas adicionales a 1300: **{semanas_adicionales}**")
with col4:
    ibl = st.number_input("Ingreso Base de Liquidación (IBL)", min_value=0.0, value=1500000.0, step=10000.0)
with col5:
    opciones_anios = list(smlmv_historico.keys())
    if current_year not in opciones_anios:
        opciones_anios.insert(0, current_year)
    opciones_anios.append("Ingresar Manualmente")
    
    anio_reconocimiento = st.selectbox("Año de Liquidación / SMLMV", options=opciones_anios)
    
    if anio_reconocimiento == "Ingresar Manualmente" or (anio_reconocimiento == current_year and current_year not in smlmv_historico):
        smlmv = st.number_input(f"Ingrese SMLMV para el año {anio_reconocimiento} ($)", min_value=1.0, value=1300000.0, step=10000.0)
    else:
        smlmv = float(smlmv_historico[anio_reconocimiento])
        st.info(f"SMLMV {anio_reconocimiento}: **${smlmv:,.0f}**")

generar_formula = st.checkbox("Explicar y detallar la fórmula decreciente en el acto administrativo")
explicacion_formula = ""

if generar_formula and smlmv > 0:
    s = ibl / smlmv
    r_base = 65.5 - (0.5 * s)
    r_base = max(55.0, min(r_base, 65.5)) 
    
    bloques_50 = int(semanas_adicionales // 50)
    porcentaje_adicional = bloques_50 * 1.5
    
    tasa_final = r_base + porcentaje_adicional
    tasa_final = min(80.0, tasa_final) 
    
    mesada = ibl * (tasa_final / 100.0)
    mesada_final = max(smlmv, mesada) 
    
    explicacion_formula = f"""Para la liquidación de la prestación económica, se procede a aplicar la fórmula decreciente consagrada en el artículo 10 de la Ley 797 de 2003 (que modificó el artículo 34 de la Ley 100 de 1993), desarrollada de la siguiente manera:

1. CÁLCULO DEL PORCENTAJE BASE (r):
La norma establece que r = 65.5 - 0.5s, donde 's' equivale al número de salarios mínimos legales mensuales vigentes (SMLMV) que representa el Ingreso Base de Liquidación (IBL).
- Año de liquidación: {anio_reconocimiento}
- SMLMV al momento del reconocimiento: ${smlmv:,.2f}
- Ingreso Base de Liquidación (IBL) calculado: ${ibl:,.2f}
- Proporción en salarios (s = IBL / SMLMV): {s:.4f}
Aplicando la fórmula (65.5 - (0.5 * {s:.4f})), se obtiene un porcentaje base inicial del: {r_base:.2f}%.

2. CÁLCULO DEL PORCENTAJE ADICIONAL POR SEMANAS:
La ley estipula un incremento del 1.5% en la tasa de reemplazo por cada 50 semanas adicionales a las primeras 1300 semanas exigidas.
- Total de semanas cotizadas acreditadas: {semanas_totales}
- Semanas adicionales (Total - 1300): {semanas_adicionales}
- Bloques completos de 50 semanas: {bloques_50}
Multiplicando los bloques ({bloques_50}) por el 1.5%, se obtiene un porcentaje adicional de: {porcentaje_adicional:.2f}%.

3. TASA DE REEMPLAZO FINAL Y MESADA PENSIONAL:
Sumando el porcentaje base ({r_base:.2f}%) y el porcentaje adicional ({porcentaje_adicional:.2f}%), se obtiene una Tasa de Reemplazo del {r_base + porcentaje_adicional:.2f}%. 
Al aplicar el tope máximo legal del 80%, la tasa definitiva a aplicar sobre el IBL es: {tasa_final:.2f}%.
- Mesada Pensional Resultante (IBL * Tasa Definitiva): ${mesada:,.2f}
- Mesada a reconocer (Aplicando garantía de pensión mínima si hubiere lugar): ${mesada_final:,.2f}
"""
    st.text_area("Vista previa del anexo de liquidación:", value=explicacion_formula, height=400)

# 5. Generación del Acto Administrativo
st.header("5. Plantilla y Generación de la Resolución")
st.write("Si no subes una plantilla, el sistema estructurará y generará una resolución oficial automáticamente.")
uploaded_file = st.file_uploader("Opcional: Suba una plantilla en Word (.docx)", type="docx")

if st.button("Generar Resolución Estructurada", type="primary"):
    # Diccionario de variables (aplica con o sin plantilla)
    reemplazos = {
        "{{NOMBRES}}": nombres.upper(),
        "{{APELLIDOS}}": apellidos.upper(),
        "{{IDENTIFICACION}}": identificacion,
        "{{NACIMIENTO}}": str(nacimiento.strftime("%d/%m/%Y")),
        "{{ANTECEDENTES}}": antecedentes,
        "{{MOTIVACIONES}}": texto_motivaciones,
        "{{FORMULA}}": explicacion_formula if generar_formula else ""
    }

    if uploaded_file is not None:
        # LÓGICA CON PLANTILLA SUBIDA
        doc = Document(uploaded_file)
        
        # Iterar sobre párrafos
        for p in doc.paragraphs:
            for key, value in reemplazos.items():
                if key in p.text:
                    p.text = p.text.replace(key, value)
                    
        # Iterar sobre tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in reemplazos.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, value)
    else:
        # LÓGICA DE GENERACIÓN AUTOMÁTICA (SIN PLANTILLA)
        doc = Document()
        doc.add_heading('REPUBLICA DE COLOMBIA', 0)
        doc.add_heading('ADMINISTRADORA COLOMBIANA DE PENSIONES - COLPENSIONES', 1)
        doc.add_paragraph('RESOLUCIÓN NÚMERO _________')
        doc.add_paragraph('RADICADO No. _________')
        doc.add_paragraph('POR LA CUAL SE RESUELVE UN TRÁMITE DE PRESTACIONES ECONÓMICAS EN EL RÉGIMEN DE PRIMA MEDIA CON PRESTACIÓN DEFINIDA')
        
        doc.add_heading('CONSIDERANDO', level=2)
        if antecedentes.strip():
            doc.add_paragraph(antecedentes)
        else:
            doc.add_paragraph('Se deja constancia de la solicitud interpuesta por el peticionario...')
            
        doc.add_heading('CONSIDERACIONES', level=2)
        if texto_motivaciones.strip():
            doc.add_paragraph(texto_motivaciones)
        else:
            doc.add_paragraph('Se incorpora el marco legal general aplicable al presente trámite...')
            
        if generar_formula:
            doc.add_heading('LIQUIDACIÓN (FÓRMULA LEY 797 DE 2003)', level=2)
            doc.add_paragraph(explicacion_formula)
            
        doc.add_heading('RESUELVE', level=2)
        doc.add_paragraph(f'ARTÍCULO PRIMERO: Reconocer y/o decidir sobre la prestación a favor de {reemplazos["{{NOMBRES}}"]} {reemplazos["{{APELLIDOS}}"]}, identificado(a) con CC No. {reemplazos["{{IDENTIFICACION}}"]}, con fecha de nacimiento {reemplazos["{{NACIMIENTO}}"]}.')
        doc.add_paragraph('ARTÍCULO SEGUNDO: Notifíquese la presente resolución, advirtiendo que proceden los recursos de ley correspondientes de conformidad con el C.P.A y de lo C.A.')
        
        doc.add_paragraph('\nCOMUNÍQUESE, NOTIFÍQUESE Y CÚMPLASE\n')
        doc.add_paragraph('[Firma Subdirector]\nSUBDIRECTOR DE DETERMINACION\nCOLPENSIONES')

    # Guardar en buffer para descarga
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    st.success("¡Acto administrativo generado exitosamente!")
    st.download_button(
        label="📥 Descargar Resolución Sustanciada (.docx)",
        data=buffer,
        file_name=f"Resolucion_{identificacion}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
